#!/usr/bin/env python3
"""Edit an existing Word document in place.

Part of the word-docx skill. Invoked via the `execute` tool:

    python skills/word-docx/scripts/edit.py \
        --file output/report.docx \
        --add-heading '{"level":2,"text":"Next Steps"}' \
        --append-paragraphs '["New paragraph here."]' \
        --append-bullets '["Step 1","Step 2"]' \
        --replace '{"find":"old term","with":"new term"}' \
        --append-blocks '[{"type":"heading","level":2,"text":"Notes"},{"type":"paragraph","text":"..."}]'

Safety rules (implement SKILL.md Key Rules):
- Minimal edits: only append/replace; never rewrite whole paragraphs.
- Find/replace reconstructs run-level text so matches split across multiple
  runs still work (a common python-docx trap).
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Edit an existing Word document (word-docx skill).",
    )
    parser.add_argument("--file", required=True, help="Document path to edit (workspace-relative).")
    parser.add_argument(
        "--append-paragraphs",
        default=None,
        help="JSON array of body paragraph strings to append.",
    )
    parser.add_argument(
        "--add-heading",
        action="append",
        default=[],
        help='JSON {"level":1,"text":"..."} for a heading to append. Repeatable.',
    )
    parser.add_argument(
        "--append-bullets",
        default=None,
        help="JSON array of bullet strings to append (List Bullet style).",
    )
    parser.add_argument(
        "--replace",
        action="append",
        default=[],
        help='JSON {"find":"old","with":"new"} for find-and-replace. Repeatable.',
    )
    parser.add_argument(
        "--append-blocks",
        default=None,
        help=(
            "JSON array of content blocks appended IN ORDER: "
            '[{"type":"heading","level":1,"text":"..."},'
            '{"type":"paragraph","text":"..."},'
            '{"type":"bullet","text":"..."},'
            '{"type":"table","headers":[...],"rows":[[...]]}]. '
            "Use this to append an interleaved heading+body section in one call."
        ),
    )
    args = parser.parse_args()

    file_path = _resolve_path(args.file)
    if not file_path.exists():
        print(f"ERROR: file not found: {file_path}", file=sys.stderr)
        return 2

    doc = Document(str(file_path))
    changes = 0

    # find-and-replace first (on existing content), then appends.
    for spec_str in args.replace:
        try:
            spec = json.loads(spec_str)
        except json.JSONDecodeError as e:
            print(f"ERROR: --replace is not valid JSON: {e}", file=sys.stderr)
            return 2
        find = str(spec.get("find", ""))
        with_ = str(spec.get("with", ""))
        changes += _replace_text(doc, find, with_)

    for h_str in args.add_heading:
        try:
            h = json.loads(h_str)
        except json.JSONDecodeError as e:
            print(f"ERROR: --add-heading is not valid JSON: {e}", file=sys.stderr)
            return 2
        level = max(1, min(9, int(h.get("level", 1))))
        doc.add_heading(str(h.get("text", "")), level=level)
        changes += 1

    if args.append_paragraphs is not None:
        try:
            paras = json.loads(args.append_paragraphs)
        except json.JSONDecodeError as e:
            print(f"ERROR: --append-paragraphs is not valid JSON: {e}", file=sys.stderr)
            return 2
        if not isinstance(paras, list):
            print("ERROR: --append-paragraphs must be a JSON array.", file=sys.stderr)
            return 2
        for p in paras:
            doc.add_paragraph(str(p))
            changes += 1

    if args.append_bullets is not None:
        try:
            bullets = json.loads(args.append_bullets)
        except json.JSONDecodeError as e:
            print(f"ERROR: --append-bullets is not valid JSON: {e}", file=sys.stderr)
            return 2
        if not isinstance(bullets, list):
            print("ERROR: --append-bullets must be a JSON array.", file=sys.stderr)
            return 2
        for b in bullets:
            doc.add_paragraph(str(b), style="List Bullet")
            changes += 1

    if args.append_blocks is not None:
        try:
            blocks = json.loads(args.append_blocks)
        except json.JSONDecodeError as e:
            print(f"ERROR: --append-blocks is not valid JSON: {e}", file=sys.stderr)
            return 2
        if not isinstance(blocks, list):
            print("ERROR: --append-blocks must be a JSON array.", file=sys.stderr)
            return 2
        for block in blocks:
            if not isinstance(block, dict):
                print(
                    f"ERROR: --append-blocks entries must be JSON objects, got {type(block).__name__}.",
                    file=sys.stderr,
                )
                return 2
            btype = block.get("type")
            if btype == "heading":
                level = max(1, min(9, int(block.get("level", 1))))
                doc.add_heading(str(block.get("text", "")), level=level)
                changes += 1
            elif btype == "paragraph":
                doc.add_paragraph(str(block.get("text", "")))
                changes += 1
            elif btype == "bullet":
                doc.add_paragraph(str(block.get("text", "")), style="List Bullet")
                changes += 1
            elif btype == "table":
                spec = block
                if not isinstance(spec, dict) or "headers" not in spec:
                    print(
                        'ERROR: --append-blocks table must be {"headers":[...],"rows":[[...]]}.',
                        file=sys.stderr,
                    )
                    return 2
                headers = [str(x) for x in spec.get("headers", [])]
                rows = spec.get("rows", [])
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    for paragraph in table.rows[0].cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                for r_idx, row in enumerate(rows, start=1):
                    for c_idx in range(len(headers)):
                        val = row[c_idx] if c_idx < len(row) else ""
                        table.rows[r_idx].cells[c_idx].text = str(val)
                changes += 1
            else:
                print(
                    f"ERROR: --append-blocks block has unknown type {btype!r} "
                    "(expected heading/paragraph/bullet/table).",
                    file=sys.stderr,
                )
                return 2

    doc.save(str(file_path))
    print(f"OK edited: {file_path}")
    print(f"  changes: {changes}")
    print(f"  paragraphs now: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
    return 0


def _replace_text(doc, find: str, replace_with: str) -> int:
    """Find-and-replace that works across runs.

    A single visible phrase is frequently split across multiple runs in
    python-docx. We rebuild each paragraph: collect all runs' text, do the
    replacement on the concatenated string, then put the result into the first
    run and clear the rest. This is the minimal-surprise approach and matches
    the SKILL.md guidance.
    """
    if not find:
        return 0
    count = 0
    for paragraph in doc.paragraphs:
        if find not in paragraph.text:
            continue
        full_text = paragraph.text
        new_text = full_text.replace(find, replace_with)
        runs = paragraph.runs
        if not runs:
            # paragraph with no runs (rare) — add one
            paragraph.add_run(new_text)
        else:
            runs[0].text = new_text
            for r in runs[1:]:
                r.text = ""
        count += 1
    # also scan inside table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if find not in paragraph.text:
                        continue
                    full_text = paragraph.text
                    new_text = full_text.replace(find, replace_with)
                    runs = paragraph.runs
                    if runs:
                        runs[0].text = new_text
                        for r in runs[1:]:
                            r.text = ""
                    else:
                        paragraph.add_run(new_text)
                    count += 1
    return count


def _resolve_path(arg: str) -> Path:
    if arg.startswith("/"):
        return Path(arg.lstrip("/"))
    return Path(arg)


if __name__ == "__main__":
    raise SystemExit(main())
