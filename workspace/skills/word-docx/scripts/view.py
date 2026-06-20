#!/usr/bin/env python3
"""Inspect (read back) a Word document for verification.

Part of the word-docx skill. Read-only — never modifies the file.

    python skills/word-docx/scripts/view.py --file output/report.docx
    python skills/word-docx/scripts/view.py --file output/report.docx --show-styles --max-paragraphs 20

NOTE: named `view.py`, not `inspect.py`, to avoid shadowing the Python
standard-library `inspect` module.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Inspect a Word document (word-docx skill, read-only).",
    )
    parser.add_argument("--file", required=True, help="Document path to inspect (workspace-relative).")
    parser.add_argument("--max-paragraphs", type=int, default=15, help="Max paragraphs to print (default 15).")
    parser.add_argument("--show-styles", action="store_true", help="Show the paragraph style name with each line.")
    args = parser.parse_args()

    file_path = _resolve_path(args.file)
    if not file_path.exists():
        print(f"ERROR: file not found: {file_path}", file=sys.stderr)
        return 2

    doc = Document(str(file_path))

    print(f"file: {file_path}")
    print(f"paragraphs: {len(doc.paragraphs)}")
    print(f"tables: {len(doc.tables)}")

    # heading outline
    headings = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style.startswith("Heading") or style == "Title":
            headings.append((style, p.text))
    print(f"headings: {len(headings)}")
    for style, text in headings:
        print(f"  [{style}] {text}")
    print()

    # first N paragraphs
    limit = max(0, args.max_paragraphs)
    printed = 0
    for p in doc.paragraphs:
        if printed >= limit:
            remaining = len(doc.paragraphs) - printed
            if remaining > 0:
                print(f"... ({remaining} more paragraphs)")
            break
        text = p.text
        if not text.strip():
            printed += 1
            continue
        if args.show_styles:
            style = p.style.name if p.style else "?"
            print(f"  [{style}] {_short(text)}")
        else:
            print(f"  {_short(text)}")
        printed += 1

    # table summary
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.columns else 0
        print()
        print(f"table[{i}]: {rows} rows x {cols} cols")
        # print first row as header preview
        if rows > 0:
            header_cells = [c.text for c in table.rows[0].cells]
            print(f"  header: {header_cells}")

    return 0


def _short(s: str) -> str:
    s = s.replace("\n", " ")
    return s if len(s) <= 100 else s[:97] + "..."


def _resolve_path(arg: str) -> Path:
    if arg.startswith("/"):
        return Path(arg.lstrip("/"))
    return Path(arg)


if __name__ == "__main__":
    raise SystemExit(main())
