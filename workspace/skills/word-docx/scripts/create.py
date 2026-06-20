#!/usr/bin/env python3
"""Create a new Word document with a title, headings, paragraphs, bullets, table.

Part of the word-docx skill. Invoked by the agent via the `execute` tool:

    python skills/word-docx/scripts/create.py \
        --out output/report.docx \
        --title "Project Report" \
        --headings '[{"level":1,"text":"Overview"}]' \
        --paragraphs '["Body text."]' \
        --bullets '["Goal one","Goal two"]' \
        --table '{"headers":["Metric","Value"],"rows":[["Users","1200"]]}'

For documents that interleave headings and body text (the common case — a
section heading followed by its paragraph, then the next heading, etc.),
use `--blocks` so everything is written in the order you give it in a single
call:

    python skills/word-docx/scripts/create.py \
        --out output/report.docx \
        --title "Project Report" \
        --blocks '[{"type":"heading","level":1,"text":"Overview"},
                   {"type":"paragraph","text":"Body text."},
                   {"type":"heading","level":2,"text":"Goals"},
                   {"type":"bullet","text":"Goal one"}]'

Design notes (implement SKILL.md Key Rules by default):
- Headings use real `Heading N` styles, not bold+font simulation.
- Bullets use the `List Bullet` style.
- Tables get explicit column widths so layout is predictable.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Create a new Word document (word-docx skill).",
    )
    parser.add_argument("--out", required=True, help="Output .docx path (workspace-relative, e.g. output/r.docx).")
    parser.add_argument("--title", default=None, help="Top-level Heading 1 title.")
    parser.add_argument(
        "--headings",
        default=None,
        help='JSON array of {"level":1,"text":"..."} (level 1-9).',
    )
    parser.add_argument(
        "--paragraphs",
        default=None,
        help="JSON array of body paragraph strings.",
    )
    parser.add_argument(
        "--bullets",
        default=None,
        help="JSON array of bullet item strings (uses List Bullet style).",
    )
    parser.add_argument(
        "--table",
        default=None,
        help='JSON object {"headers":[...],"rows":[[...],...]} for one table.',
    )
    parser.add_argument(
        "--blocks",
        default=None,
        help=(
            'JSON array of content blocks written IN ORDER: '
            '[{"type":"heading","level":1,"text":"..."},'
            '{"type":"paragraph","text":"..."},'
            '{"type":"bullet","text":"..."},'
            '{"type":"table","headers":[...],"rows":[[...]]}]. '
            "Use this for interleaved heading+body documents. "
            "When set, --headings/--paragraphs/--bullets/--table are ignored."
        ),
    )
    args = parser.parse_args()

    out_path = _resolve_out_path(args.out)

    doc = Document()

    n_headings = 0
    n_paragraphs = 0
    n_bullets = 0
    n_table_rows = 0
    n_blocks = 0

    if args.title:
        doc.add_heading(args.title, level=1)

    if args.blocks:
        # Warn if the legacy batch args were also supplied — they are ignored.
        ignored = [
            name
            for name, val in (
                ("--headings", args.headings),
                ("--paragraphs", args.paragraphs),
                ("--bullets", args.bullets),
                ("--table", args.table),
            )
            if val
        ]
        if ignored:
            print(
                "WARNING: --blocks takes precedence; ignoring "
                + ", ".join(ignored),
                file=sys.stderr,
            )
        blocks = _parse_json_arg(args.blocks, "--blocks", expect_list=True)
        if blocks is None:
            return 2
        for block in blocks:
            added = _add_block(doc, block, "--blocks")
            if added is None:
                return 2
            btype = added[0]
            count = added[1]
            n_blocks += 1
            if btype == "heading":
                n_headings += count
            elif btype == "paragraph":
                n_paragraphs += count
            elif btype == "bullet":
                n_bullets += count
            elif btype == "table":
                n_table_rows += count
    else:
        if args.headings:
            headings = _parse_json_arg(args.headings, "--headings", expect_list=True)
            if headings is None:
                return 2
            for h in headings:
                level = int(h.get("level", 1))
                level = max(1, min(9, level))
                text = str(h.get("text", ""))
                doc.add_heading(text, level=level)
                n_headings += 1

        if args.paragraphs:
            paras = _parse_json_arg(args.paragraphs, "--paragraphs", expect_list=True)
            if paras is None:
                return 2
            for p in paras:
                doc.add_paragraph(str(p))
                n_paragraphs += 1

        if args.bullets:
            bullets = _parse_json_arg(args.bullets, "--bullets", expect_list=True)
            if bullets is None:
                return 2
            for b in bullets:
                doc.add_paragraph(str(b), style="List Bullet")
                n_bullets += 1

        if args.table:
            tbl = _parse_json_arg(args.table, "--table", expect_list=False)
            if tbl is None:
                return 2
            n_table_rows = _add_table_from_spec(doc, tbl)
            if n_table_rows is None:
                return 2

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    print(f"OK created: {out_path}")
    print(f"  paragraphs: {len(doc.paragraphs)}")
    print(f"    - headings: {n_headings}")
    print(f"    - body paragraphs: {n_paragraphs}")
    print(f"    - bullets: {n_bullets}")
    print(f"  tables: {len(doc.tables)} (rows added: {n_table_rows})")
    if args.blocks:
        print(f"  blocks: {n_blocks}")
    return 0


def _parse_json_arg(raw: str, name: str, expect_list: bool):
    """Parse a JSON CLI arg. Returns parsed value or None on error."""
    try:
        value = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"ERROR: {name} is not valid JSON: {e}", file=sys.stderr)
        return None
    if expect_list and not isinstance(value, list):
        print(f"ERROR: {name} must be a JSON array.", file=sys.stderr)
        return None
    return value


def _add_block(doc, block, arg_name: str):
    """Write a single content block to `doc` in order.

    Returns (type, count) on success, or None on error. `count` is the number
    of elements added by this block (1 for heading/paragraph/bullet,
    number of data rows for a table).
    """
    if not isinstance(block, dict):
        print(
            f"ERROR: {arg_name} entries must be JSON objects, got {type(block).__name__}.",
            file=sys.stderr,
        )
        return None
    btype = block.get("type")
    if btype == "heading":
        level = max(1, min(9, int(block.get("level", 1))))
        text = str(block.get("text", ""))
        doc.add_heading(text, level=level)
        return ("heading", 1)
    if btype == "paragraph":
        doc.add_paragraph(str(block.get("text", "")))
        return ("paragraph", 1)
    if btype == "bullet":
        doc.add_paragraph(str(block.get("text", "")), style="List Bullet")
        return ("bullet", 1)
    if btype == "table":
        rows = _add_table_from_spec(doc, block)
        if rows is None:
            return None
        return ("table", rows)
    print(
        f"ERROR: {arg_name} block has unknown type {btype!r} "
        "(expected heading/paragraph/bullet/table).",
        file=sys.stderr,
    )
    return None


def _add_table_from_spec(doc, spec) -> int | None:
    """Add one table from a {"headers":[...],"rows":[[...]]} spec.

    Returns the number of data rows added, or None on error.
    """
    if not isinstance(spec, dict) or "headers" not in spec:
        print(
            'ERROR: table spec must be {"headers":[...],"rows":[[...]]}.',
            file=sys.stderr,
        )
        return None
    headers = [str(x) for x in spec.get("headers", [])]
    rows = spec.get("rows", [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # data rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(len(headers)):
            val = row[c_idx] if c_idx < len(row) else ""
            table.rows[r_idx].cells[c_idx].text = str(val)
    # explicit column widths so layout is predictable
    if headers:
        width = Inches(6.0 / len(headers))
        for col in table.columns:
            for cell in col.cells:
                cell.width = width
    return len(rows)


def _resolve_out_path(out_arg: str) -> Path:
    if out_arg.startswith("/"):
        return Path(out_arg.lstrip("/"))
    return Path(out_arg)


if __name__ == "__main__":
    raise SystemExit(main())
